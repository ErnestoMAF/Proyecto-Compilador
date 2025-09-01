from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import qn
import os
import subprocess
import platform

class ModificadorWord:
    def __init__(self, archivo=None):
        """
        Inicializa el documento. Si no se proporciona un archivo, crea uno nuevo.
        Si se proporciona un archivo que existe, lo carga. Si no existe, crea uno nuevo.
        :param archivo: Ruta al archivo de Word (opcional).
        """
        if archivo:
            if os.path.exists(archivo):
                try:
                    self.doc = Document(archivo)
                    print(f"Documento '{archivo}' cargado exitosamente")
                except Exception as e:
                    print(f"Error al abrir el archivo: {e}")
                    print("Creando nuevo documento...")
                    self.doc = Document()
            else:
                print(f"El archivo '{archivo}' no existe. Creando nuevo documento...")
                self.doc = Document()
        else:
            self.doc = Document()
            print("Nuevo documento creado")
    
    def cargar_documento_existente(self, archivo):
        """
        Carga un documento existente para modificarlo.
        :param archivo: Ruta al archivo de Word existente.
        :return: True si se cargó exitosamente, False si hubo error.
        """
        try:
            if os.path.exists(archivo):
                self.doc = Document(archivo)
                print(f"Documento '{archivo}' cargado exitosamente")
                return True
            else:
                print(f"Error: El archivo '{archivo}' no existe")
                return False
        except Exception as e:
            print(f"Error al cargar el documento: {e}")
            return False
    
    def guardar(self, nombre_archivo):
        """Guarda el documento con el nombre especificado."""
        self.doc.save(nombre_archivo)
        print(f"Documento guardado como '{nombre_archivo}'")
    
    def reemplazar_texto(self, texto_antiguo, texto_nuevo):
        """
        Reemplaza texto en todo el documento (incluyendo tablas).
        :param texto_antiguo: Texto a buscar.
        :param texto_nuevo: Texto de reemplazo.
        """
        for paragraph in self.doc.paragraphs:
            if texto_antiguo in paragraph.text:
                paragraph.text = paragraph.text.replace(texto_antiguo, texto_nuevo)
        
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if texto_antiguo in cell.text:
                        cell.text = cell.text.replace(texto_antiguo, texto_nuevo)
    
    def agregar_texto(self, texto, estilo='Normal'):
        """Agrega texto al final del documento."""
        self.doc.add_paragraph(texto, style=estilo)
    
    def agregar_texto_pagina_especifica(self, texto, pagina_aproximada, estilo='Normal'):
        """
        **ADVERTENCIA**: Word no expone la paginación via programación.
        Esta función agrega texto en una posición aproximada (número de párrafo).
        :param texto: Texto a agregar.
        :param pagina_aproximada: Número de página aproximado (basado en conteo de párrafos).
        :param estilo: Estilo de párrafo (por defecto 'Normal').
        """
        # Estimación muy básica: asume 3 párrafos por página
        posicion_aproximada = pagina_aproximada * 3
        num_parrafos = len(self.doc.paragraphs)
        
        if posicion_aproximada >= num_parrafos:
            # Si la posición está más allá del número de párrafos, agregar al final
            self.doc.add_paragraph(texto, style=estilo)
        else:
            # Insertar antes del párrafo en la posición aproximada
            p = self.doc.paragraphs[posicion_aproximada].insert_paragraph_before(texto, style=estilo)
    
    def formatear_texto(self, texto_buscar, negrita=False, color_rgb=None, 
                       tamaño_fuente=None, tipo_fuente=None, alineacion=None):
        """
        Aplica formato al texto que coincida con 'texto_buscar'.
        :param texto_buscar: Texto a formatear.
        :param negrita: True para negrita.
        :param color_rgb: Tupla (R, G, B) para el color.
        :param tamaño_fuente: Tamaño de fuente en puntos.
        :param tipo_fuente: Nombre de la fuente (ej: 'Arial', 'Times New Roman')
        :param alineacion: Tipo de alineación ('justificado', 'izquierda', 'derecha', 'centro')
        """
        for paragraph in self.doc.paragraphs:
            if texto_buscar in paragraph.text:
                
                # Aplicar alineación al párrafo completo
                if alineacion:
                    if alineacion.lower() == 'justificado':
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    elif alineacion.lower() == 'izquierda':
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    elif alineacion.lower() == 'derecha':
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    elif alineacion.lower() == 'centro':
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Aplicar formato a cada run (fragmento de texto)
                for run in paragraph.runs:
                    if texto_buscar in run.text:
                        run.bold = negrita
                        
                        if color_rgb:
                            run.font.color.rgb = RGBColor(*color_rgb)
                        
                        if tamaño_fuente:
                            run.font.size = Pt(tamaño_fuente)
                        
                        if tipo_fuente:
                            run.font.name = tipo_fuente
                            # Para asegurar compatibilidad con Word
                            r = run._element
                            r.rPr.rFonts.set(qn('w:eastAsia'), tipo_fuente)
    
    def formatear_parrafo_completo(self, texto_buscar, alineacion=None):
        """
        Aplica formato de alineación a párrafos completos que contengan el texto.
        :param texto_buscar: Texto a buscar en el párrafo.
        :param alineacion: Tipo de alineación ('justificado', 'izquierda', 'derecha', 'centro')
        """
        for paragraph in self.doc.paragraphs:
            if texto_buscar in paragraph.text and alineacion:
                if alineacion.lower() == 'justificado':
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                elif alineacion.lower() == 'izquierda':
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                elif alineacion.lower() == 'derecha':
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                elif alineacion.lower() == 'centro':
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def escribir_despues_de_texto(self, texto_buscar, texto_agregar, ubicacion='inmediata', mismo_parrafo=True):
        """
        Escribe texto después de encontrar el texto buscado.
        
        :param texto_buscar: Texto a buscar en el documento
        :param texto_agregar: Texto a agregar después del texto encontrado
        :param ubicacion: 'inmediata' (justo después) o 'final_parrafo' (al final del párrafo)
        :param mismo_parrafo: Si True, agrega en el mismo párrafo. Si False, crea nuevo párrafo
        """
        texto_encontrado = False
        
        # Buscar en párrafos normales
        for paragraph in self.doc.paragraphs:
            if texto_buscar in paragraph.text:
                texto_encontrado = True
                
                if ubicacion == 'inmediata':
                    # Reemplazar el texto buscado por: texto_buscado + texto_agregar
                    nuevo_texto = paragraph.text.replace(texto_buscar, texto_buscar + texto_agregar)
                    paragraph.text = nuevo_texto
                
                elif ubicacion == 'final_parrafo':
                    if mismo_parrafo:
                        # Agregar al final del mismo párrafo
                        paragraph.text += texto_agregar
                    else:
                        # Crear nuevo párrafo después del actual
                        nuevo_parrafo = paragraph.insert_paragraph_before(texto_agregar)
                        # Mover el nuevo párrafo después del actual
                        paragraph._p.addnext(nuevo_parrafo._p)
        
        # Buscar en tablas
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if texto_buscar in paragraph.text:
                            texto_encontrado = True
                            
                            if ubicacion == 'inmediata':
                                nuevo_texto = paragraph.text.replace(texto_buscar, texto_buscar + texto_agregar)
                                paragraph.text = nuevo_texto
                            
                            elif ubicacion == 'final_parrafo':
                                if mismo_parrafo:
                                    paragraph.text += texto_agregar
                                else:
                                    # En tablas, simplemente agregamos al final
                                    paragraph.text += texto_agregar
        
        if not texto_encontrado:
            print(f"Texto '{texto_buscar}' no encontrado en el documento")
    
    def escribir_antes_de_texto(self, texto_buscar, texto_agregar):
        """
        Escribe texto antes de encontrar el texto buscado.
        """
        texto_encontrado = False
        
        for paragraph in self.doc.paragraphs:
            if texto_buscar in paragraph.text:
                texto_encontrado = True
                nuevo_texto = paragraph.text.replace(texto_buscar, texto_agregar + texto_buscar)
                paragraph.text = nuevo_texto
        
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if texto_buscar in paragraph.text:
                            texto_encontrado = True
                            nuevo_texto = paragraph.text.replace(texto_buscar, texto_agregar + texto_buscar)
                            paragraph.text = nuevo_texto
        
        if not texto_encontrado:
            print(f"Texto '{texto_buscar}' no encontrado en el documento")
    
    def obtener_posicion_texto(self, texto_buscar):
        """
        Devuelve información sobre la posición del texto encontrado.
        Retorna: (parrafo_index, posicion_en_parrafo, texto_del_parrafo)
        """
        for i, paragraph in enumerate(self.doc.paragraphs):
            if texto_buscar in paragraph.text:
                posicion = paragraph.text.find(texto_buscar)
                return i, posicion, paragraph.text
        
        for table_index, table in enumerate(self.doc.tables):
            for row_index, row in enumerate(table.rows):
                for cell_index, cell in enumerate(row.cells):
                    for para_index, paragraph in enumerate(cell.paragraphs):
                        if texto_buscar in paragraph.text:
                            posicion = paragraph.text.find(texto_buscar)
                            return (f"tabla_{table_index}_fila_{row_index}_celda_{cell_index}_parrafo_{para_index}", 
                                    posicion, paragraph.text)
        
        return None, None, None
    
    def abrir_con_aplicacion(self, nombre_archivo=None):
        """
        Abre el documento con la aplicación predeterminada del sistema.
        :param nombre_archivo: Nombre del archivo a abrir (opcional, usa el documento actual si no se especifica)
        """
        if nombre_archivo is None:
            # Si no se especifica archivo, usa el último guardado o pide uno
            print("Error: Debe especificar un nombre de archivo")
            return False
        
        try:
            # Verificar si el archivo existe
            if not os.path.exists(nombre_archivo):
                print(f"Error: El archivo '{nombre_archivo}' no existe")
                return False
            
            # Abrir con la aplicación predeterminada según el sistema operativo
            sistema = platform.system()
            
            if sistema == "Windows":
                os.startfile(nombre_archivo)
                print(f"Abriendo '{nombre_archivo}' con la aplicación predeterminada de Windows")
            
            elif sistema == "Linux":
                subprocess.run(["xdg-open", nombre_archivo])
                print(f"Abriendo '{nombre_archivo}' con la aplicación predeterminada de Linux")
            
            else:
                print(f"Sistema operativo no compatible: {sistema}")
                return False
            
            return True
            
        except Exception as e:
            print(f"Error al abrir el archivo: {e}")
            return False

# Ejemplo de uso completo
if __name__ == "__main__":
    nombre_archivo = "documento_modificado.docx"
    
    # Verificar si el archivo existe para cargarlo o crear uno nuevo
    if os.path.exists(nombre_archivo):
        print(f"El archivo '{nombre_archivo}' existe. Cargándolo...")
        doc = ModificadorWord(nombre_archivo)
    else:
        print(f"El archivo '{nombre_archivo}' no existe. Creando nuevo documento...")
        doc = ModificadorWord()
        
        # Agregar contenido inicial solo si es un documento nuevo
        doc.agregar_texto("Documento de Prueba Completo")
        doc.agregar_texto("Estimado {cliente},")
        doc.agregar_texto("Le informamos sobre los RESULTADOS de su solicitud.")
        doc.agregar_texto("Este es un párrafo de ejemplo con texto importante.")
    
    # Realizar modificaciones (tanto para documentos nuevos como existentes)
    
    # 1. Reemplazar texto
    doc.reemplazar_texto("{cliente}", "Pepe")
    
    # 2. Formatear texto
    doc.formatear_texto(
        texto_buscar="RESULTADOS",
        negrita=True,
        color_rgb=(255, 0, 0),
        tamaño_fuente=14,
        tipo_fuente="Lexend",
        alineacion="centro"
    )
    
    # 3. Escribir después de texto (inmediatamente)
    doc.escribir_despues_de_texto(
        texto_buscar="RESULTADOS",
        texto_agregar=" que fueron exitosos",
        ubicacion='inmediata'
    )
    
    # 4. Escribir al final del párrafo
    doc.escribir_despues_de_texto(
        texto_buscar="ejemplo",
        texto_agregar=" [Texto agregado al final]",
        ubicacion='final_parrafo'
    )
    
    # 5. Escribir antes de texto
    doc.escribir_antes_de_texto(
        texto_buscar="importante",
        texto_agregar="muy "
    )
    
    # 6. Obtener posición de texto
    posicion = doc.obtener_posicion_texto("RESULTADOS")
    print(f"Posición del texto: {posicion}")
    
    # 7. Formatear párrafo completo
    doc.formatear_parrafo_completo(
        texto_buscar="párrafo de ejemplo",
        alineacion="justificado"
    )
    
    # 8. Agregar texto en posición aproximada (página 2)
    # Nota: Si el documento tiene menos de 6 párrafos, se agregará al final.
    doc.agregar_texto_pagina_especifica(
        texto="Texto agregado en página aproximada 2",
        pagina_aproximada=2
    )
    
    # Guardar el documento
    doc.guardar(nombre_archivo)

    # Abrir el documento con la aplicación predeterminada
    doc.abrir_con_aplicacion(nombre_archivo)